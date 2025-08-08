import logging
import os
import fitz  # PyMuPDF
import base64
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass
from PIL import Image
import html2text
import chardet
import io
import hashlib

# Import langchain components
from langchain.schema.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownTextSplitter

# PyMuPDF4LLM for markdown extraction
try:
    import pymupdf4llm
    PYMUPDF4LLM_AVAILABLE = True
except ImportError:
    PYMUPDF4LLM_AVAILABLE = False
    print("⚠️ Warning: pymupdf4llm not available. Install with: pip install pymupdf4llm")

# Configure logging first
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Document processing
# Import optional dependencies with fallbacks
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. DOCX processing will be limited.")
    DocxDocument = None
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    logger.warning("pandas not available. Spreadsheet processing will be limited.")
    pd = None
    PANDAS_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    logger.warning("striprtf not available. RTF processing will be limited.")
    rtf_to_text = None
    RTF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    logger.warning("BeautifulSoup not available. HTML processing will be limited.")
    BeautifulSoup = None
    BS4_AVAILABLE = False

try:
    import markdown
    from markdown.extensions import codehilite, tables
    MARKDOWN_AVAILABLE = True
except ImportError:
    markdown = None
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown not installed. Markdown support will be limited.")

# Ollama imports with fallbacks
try:
    from langchain_ollama import OllamaLLM as Ollama
except ImportError:
    try:
        from langchain_community.llms import Ollama
    except ImportError:
        try:
            from langchain.llms import Ollama
        except ImportError:
            logger.warning("Ollama not available. Please install langchain-ollama or langchain-community")
            Ollama = None

# Pipeline import for HuggingFace models
try:
    from transformers import pipeline
    import torch
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    logger.warning("transformers not available. Some VLM models will not work.")
    pipeline = None
    torch = None
    TRANSFORMERS_AVAILABLE = False

@dataclass
class ExtractedElement:
    """Container for extracted document elements."""
    element_type: str  # 'text', 'image', 'table', 'graph'
    content: Any  # Text content, image bytes, HTML table, etc.
    metadata: Dict[str, Any]
    page_number: int
    element_id: str
    bbox: Optional[Tuple[float, float, float, float]] = None  # Bounding box coordinates

class MultimodalDocumentProcessor:
    """
    Multimodal document processor that extracts and processes different types of content
    from documents (PDF, DOCX) according to specific processing strategies.
    """
    
    def __init__(self, 
                 vlm_model_name: str = "gemma3:27b",
                 llm_model_name: str = "gemma3:27b",
                 chunk_size: int = 800,
                 chunk_overlap: int = 80):
        """
        Initialize the multimodal document processor.
        
        Args:
            vlm_model_name: Vision Language Model for image description (default: gemma3:27b)
            llm_model_name: Language Model for table descriptions (default: gemma3:27b)
            chunk_size: Text chunk size for splitting
            chunk_overlap: Overlap between text chunks
        """
        # Ensure models are never None - use gemma3:27b as default
        self.vlm_model_name = vlm_model_name if vlm_model_name is not None else "gemma3:27b"
        self.llm_model_name = llm_model_name if llm_model_name is not None else "gemma3:27b"
        
        # Log warnings if None values were passed
        if vlm_model_name is None:
            logger.warning("VLM model was None, using default: gemma3:27b")
        if llm_model_name is None:
            logger.warning("LLM model was None, using default: gemma3:27b")
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        logger.info(f"Initializing document processor with VLM: {self.vlm_model_name}, LLM: {self.llm_model_name}")
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )
        
        # Initialize VLM for image description
        self._init_vlm()
        
        # Initialize LLM for table descriptions (optional)
        self._init_llm()
        
        # Initialize HTML converter for tables
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = True
        self.html_converter.ignore_images = True
        
    def _init_vlm(self):
        """Initialize Vision Language Model (VLM)."""
        if not self.vlm_model_name:
            logger.warning("No VLM model specified")
            return

        try:
            logger.info(f"Loading VLM model: {self.vlm_model_name}")
            
            # Check if this is an Ollama model (contains colon)
            if ":" in self.vlm_model_name or self.vlm_model_name.startswith("gemma3") or self.vlm_model_name.startswith("llava") or self.vlm_model_name.startswith("llama3.2-vision"):
                # Use Ollama for models like gemma3:4b, llava:7b, etc.
                logger.info(f"Using Ollama for VLM: {self.vlm_model_name}")
                
                # Try different import options for Ollama
                try:
                    from langchain_ollama import OllamaLLM as Ollama
                except ImportError:
                    try:
                        from langchain_community.llms import Ollama
                    except ImportError:
                        try:
                            from langchain.llms import Ollama
                        except ImportError:
                            logger.error("Could not import Ollama. Please install langchain-community or langchain-ollama")
                            return

                self.vlm_model = Ollama(model=self.vlm_model_name)
                self.vlm_processor = None  # Not needed for Ollama
                self.use_ollama_vlm = True
                logger.info(f"Ollama VLM model loaded: {self.vlm_model_name}")
                
            else:
                # Use HuggingFace transformers for models like microsoft/git-base-coco
                logger.info(f"Using HuggingFace for VLM: {self.vlm_model_name}")
                from transformers import AutoProcessor, AutoModelForCausalLM
                
                self.vlm_processor = AutoProcessor.from_pretrained(self.vlm_model_name)
                self.vlm_model = AutoModelForCausalLM.from_pretrained(self.vlm_model_name)
                self.vlm_model.to(self.device)
                self.use_ollama_vlm = False
                logger.info(f"HuggingFace VLM model loaded: {self.vlm_model_name}")
                
        except Exception as e:
            logger.error(f"Failed to load VLM model: {e}")
            self.vlm_model = None
    
    def _init_llm(self):
        """Initialize Language Model (LLM) for table descriptions."""
        # Ensure we have a valid LLM model name
        if not self.llm_model_name or self.llm_model_name.strip() == "":
            self.llm_model_name = "gemma3:27b"
            logger.warning("LLM model was empty/None, using default: gemma3:27b")
        
        try:
            logger.info(f"Loading LLM model: {self.llm_model_name}")
            
            # Check if this is an Ollama model (contains colon or known Ollama models)
            if (":" in self.llm_model_name or 
                self.llm_model_name.startswith(("gemma", "llama", "mistral", "codellama"))):
                
                # Use Ollama for models like gemma3:27b, llama3.2:3b, etc.
                logger.info(f"Using Ollama for LLM: {self.llm_model_name}")
                
                # Try different import options for Ollama
                try:
                    from langchain_ollama import OllamaLLM as Ollama
                except ImportError:
                    try:
                        from langchain_community.llms import Ollama
                    except ImportError:
                        try:
                            from langchain.llms import Ollama
                        except ImportError:
                            logger.error("Ollama not available. Please install langchain-ollama or langchain-community")
                            self.llm_model = None
                            return
                
                self.llm_model = Ollama(model=self.llm_model_name)
                logger.info(f"Ollama LLM {self.llm_model_name} loaded successfully")
                
            else:
                # Use HuggingFace for standard models
                logger.info(f"Using HuggingFace for LLM: {self.llm_model_name}")
                
                # For now, use a simple text generation pipeline
                # This can be extended to support more sophisticated models
                try:
                    self.llm_model = pipeline(
                        "text-generation",
                        model=self.llm_model_name,
                        device=0 if torch.cuda.is_available() else -1,
                        max_new_tokens=200,
                        temperature=0.7
                    )
                    logger.info(f"HuggingFace LLM {self.llm_model_name} loaded successfully")
                except Exception as e:
                    logger.warning(f"Failed to load HuggingFace LLM: {e}")
                    self.llm_model = None
                    
        except Exception as e:
            logger.error(f"Failed to load LLM model: {e}")
            self.llm_model = None
    
    def _generate_element_id(self, content: Any, page_num: int, element_type: str) -> str:
        """Generate unique ID for extracted elements."""
        content_hash = hashlib.md5(str(content).encode()).hexdigest()[:8]
        return f"{element_type}_{page_num}_{content_hash}"
    
    def _describe_image_with_vlm(self, image: Image.Image, surrounding_text: str = "") -> str:
        """Generate description for image using VLM."""
        if not self.vlm_model:
            return "Image description unavailable (VLM not loaded)"

        try:
            if hasattr(self, 'use_ollama_vlm') and self.use_ollama_vlm:
                # Use Ollama for multimodal description
                return self._describe_image_with_ollama(image, surrounding_text)
            else:
                # Use HuggingFace transformers
                return self._describe_image_with_huggingface(image, surrounding_text)
                
        except Exception as e:
            logger.error(f"Error generating image description: {e}")
            return f"Error generating image description: {str(e)}"

    def _describe_image_with_ollama(self, image: Image.Image, surrounding_text: str = "") -> str:
        """Generate description using Ollama VLM."""
        try:
            import tempfile
            import os
            
            # Save image temporarily for Ollama
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                image.save(tmp_file.name, format='PNG')
                temp_image_path = tmp_file.name
            
            try:
                # Create prompt for Ollama multimodal models
                context_text = f"Context: {surrounding_text}" if surrounding_text else ""
                prompt = f"""Please describe this image in detail. Focus on the main subjects, activities, text content, charts, graphs, or any other important visual elements.

{context_text}

What do you see in this image?"""

                # For Ollama multimodal, we need to check if it supports direct API calls
                # If using ollama Python library directly:
                try:
                    import ollama
                    response = ollama.chat(
                        model=self.vlm_model_name,
                        messages=[{
                            'role': 'user',
                            'content': prompt,
                            'images': [temp_image_path]
                        }]
                    )
                    description = response['message']['content']
                except ImportError:
                    # Fallback to langchain if ollama library not available
                    # Note: This might not work for multimodal, depends on langchain implementation
                    logger.warning("ollama library not found, falling back to langchain (multimodal support limited)")
                    description = self.vlm_model.invoke(prompt)
                
                return description.strip() if description else "Could not generate image description"
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_image_path):
                    os.unlink(temp_image_path)
            
        except Exception as e:
            logger.error(f"Error with Ollama VLM: {e}")
            return f"Image description unavailable (Ollama error: {str(e)})"

    def _describe_image_with_huggingface(self, image: Image.Image, surrounding_text: str = "") -> str:
        """Generate description using HuggingFace transformers."""
        try:
            prompt_text = f"Describe this image in detail. Context: {surrounding_text}" if surrounding_text else "Describe this image in detail."
            
            # Different models have different input formats
            if "git" in self.vlm_model_name.lower():
                # For GIT models
                inputs = self.vlm_processor(images=image, text=prompt_text, return_tensors="pt").to(self.device)
                with torch.no_grad():
                    generated_ids = self.vlm_model.generate(**inputs, max_length=150, num_beams=4, early_stopping=True)
                description = self.vlm_processor.decode(generated_ids[0], skip_special_tokens=True)
            elif "blip" in self.vlm_model_name.lower():
                # For BLIP models
                inputs = self.vlm_processor(image, prompt_text, return_tensors="pt").to(self.device)
                with torch.no_grad():
                    generated_ids = self.vlm_model.generate(**inputs, max_length=150)
                description = self.vlm_processor.decode(generated_ids[0], skip_special_tokens=True)
            else:
                # Generic approach
                inputs = self.vlm_processor(images=image, text=prompt_text, return_tensors="pt").to(self.device)
                with torch.no_grad():
                    generated_ids = self.vlm_model.generate(**inputs, max_length=150, num_beams=4, early_stopping=True)
                description = self.vlm_processor.decode(generated_ids[0], skip_special_tokens=True)
            
            # Clean up the description
            if prompt_text in description:
                description = description.replace(prompt_text, "").strip()
            
            return description if description else "Could not generate image description"
            
        except Exception as e:
            logger.error(f"Error with HuggingFace VLM: {e}")
            return f"Image description unavailable (HF error: {str(e)})"
    
    def _extract_table_description(self, table_html: str, surrounding_text: str = "") -> str:
        """
        Generate description for a table using LLM.
        
        Args:
            table_html: HTML representation of the table
            surrounding_text: Text context around the table
            
        Returns:
            Description of the table
        """
        try:
            # Convert HTML to readable text
            table_text = self.html_converter.handle(table_html)
            
            # Create prompt for table description
            prompt = f"""
            Analyze the following table and provide a comprehensive description including:
            1. The main purpose and content of the table
            2. Key columns and their meanings
            
            Context: {surrounding_text}
            
            Table:
            {table_text}
            
            Description:
            """

            try: 
                import ollama
                response = ollama.chat(
                    model=self.llm_model_name, # Use self.llm_model_name here
                    messages=[{
                        'role': 'user',
                        'content': prompt,
                    }]
                )
                description = response['message']['content']
            except ImportError:
                # Fallback to langchain if ollama library not available
                # Note: This might not work for multimodal, depends on langchain implementation
                logger.warning("ollama library not found, falling back to langchain (multimodal support limited)")
                description = self.llm_model.invoke(prompt) # Use self.llm_model here
            
            # Use a simple text generation approach (can be replaced with more sophisticated LLM)
            
            # For now, return a structured description
            return description.strip()
            
        except Exception as e:
            logger.error(f"Error generating table description: {e}")
            return f"Table description error: {str(e)}"
    
    def _save_image(self, image: Image.Image, save_dir: str, filename: str) -> str:
        """Save image to disk and return path."""
        os.makedirs(save_dir, exist_ok=True)
        image_path = os.path.join(save_dir, filename)
        image.save(image_path)
        return image_path
    
    def _save_page_as_image(self, page, save_dir: str, filename: str) -> str:
        """Save PDF page as image."""
        os.makedirs(save_dir, exist_ok=True)
        # Render page as image
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
        img_data = pix.tobytes("png")
        
        # Save image
        image_path = os.path.join(save_dir, filename)
        with open(image_path, "wb") as f:
            f.write(img_data)
        
        return image_path
    
    def process_pdf(self, pdf_path: str, output_dir: str = "extracted_content") -> List[ExtractedElement]:
        """
        Process PDF file and extract different types of content.
        Now uses markdown extraction strategy for better text processing.
        
        Args:
            pdf_path: Path to PDF file
            output_dir: Directory to save extracted images
            
        Returns:
            List of extracted elements
        """
        elements = []
        doc_name = Path(pdf_path).stem
        image_dir = os.path.join(output_dir, "images", doc_name)
        page_image_dir = os.path.join(output_dir, "pages", doc_name)
        
        try:
            doc = fitz.open(pdf_path)
            
            # STEP 1: Extract text using markdown strategy
            # This provides better structure-aware chunking
            logger.info("🔤 Extracting text using markdown strategy...")
            text_elements = self._extract_text_with_markdown_strategy(pdf_path)
            elements.extend(text_elements)
            
            # STEP 2: Process each page for images, tables, and page-level content
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Save entire page as image for multimodal retrieval
                page_image_filename = f"page_{page_num + 1}.png"
                page_image_path = self._save_page_as_image(page, page_image_dir, page_image_filename)
                
                # Create page image element
                page_element = ExtractedElement(
                    element_type="page_image",
                    content=page_image_path,
                    metadata={
                        "source": pdf_path,
                        "page": page_num + 1,
                        "document_name": doc_name,
                        "image_path": page_image_path,
                        "content_type": "page_image"
                    },
                    page_number=page_num + 1,
                    element_id=self._generate_element_id(page_image_path, page_num + 1, "page_image")
                )
                elements.append(page_element)
                
                # Get page text for context (not for chunking, just for surrounding context)
                page_text = page.get_text()
                
                # Extract images from this page
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # Ensure it's not CMYK
                            img_data = pix.tobytes("png")
                            image = Image.open(io.BytesIO(img_data))
                            
                            # Save image
                            img_filename = f"img_{page_num + 1}_{img_index}.png"
                            img_path = self._save_image(image, image_dir, img_filename)
                            
                            # Generate description using surrounding page text as context
                            surrounding_text = page_text if page_text else ""
                            description = self._describe_image_with_vlm(image, surrounding_text)
                            
                            # Determine if it's a graph or regular image
                            is_graph = any(keyword in description.lower() 
                                         for keyword in ['chart', 'graph', 'plot', 'diagram', 'visualization'])
                            element_type = "graph" if is_graph else "image"
                            
                            img_element = ExtractedElement(
                                element_type=element_type,
                                content=img_path,
                                metadata={
                                    "source": pdf_path,
                                    "page": page_num + 1,
                                    "document_name": doc_name,
                                    "image_path": img_path,
                                    "description": description,
                                    "surrounding_text": surrounding_text,
                                    "content_type": "image"
                                },
                                page_number=page_num + 1,
                                element_id=self._generate_element_id(img_path, page_num + 1, element_type)
                            )
                            elements.append(img_element)
                        
                        pix = None
                        
                    except Exception as e:
                        logger.error(f"Error processing image {img_index} on page {page_num + 1}: {e}")
                        continue
                
                # Extract tables from this page
                tables = page.find_tables()
                for table_index, table in enumerate(tables):
                    try:
                        # Extract table as HTML
                        table_html = table.to_pandas().to_html(index=False, escape=False)
                        
                        # Generate description using page text as context
                        surrounding_text = page_text if page_text else ""
                        description = self._describe_table_with_llm(table_html, surrounding_text)
                        
                        table_element = ExtractedElement(
                            element_type="table",
                            content=description,
                            metadata={
                                "source": pdf_path,
                                "page": page_num + 1,
                                "document_name": doc_name,
                                "table_html": table_html,
                                "description": description,
                                "surrounding_text": surrounding_text,
                                "content_type": "table"
                            },
                            page_number=page_num + 1,
                            element_id=self._generate_element_id(table_html, page_num + 1, "table")
                        )
                        elements.append(table_element)
                        
                    except Exception as e:
                        logger.error(f"Error processing table {table_index} on page {page_num + 1}: {e}")
                        continue
            
            doc.close()
            
            logger.info(f"✅ PDF processing completed: {len(elements)} elements extracted")
            logger.info(f"   📄 Text elements: {len([e for e in elements if e.element_type == 'text'])}")
            logger.info(f"   🖼️ Images: {len([e for e in elements if e.element_type == 'image'])}")
            logger.info(f"   📊 Graphs: {len([e for e in elements if e.element_type == 'graph'])}")
            logger.info(f"   📋 Tables: {len([e for e in elements if e.element_type == 'table'])}")
            logger.info(f"   📑 Page images: {len([e for e in elements if e.element_type == 'page_image'])}")
            
            return elements
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return []
    
    def process_docx(self, docx_path: str, output_dir: str = "extracted_content") -> List[ExtractedElement]:
        """
        Process DOCX file and extract different types of content.
        
        Args:
            docx_path: Path to DOCX file
            output_dir: Directory to save extracted images
            
        Returns:
            List of extracted elements
        """
        elements = []
        doc_name = Path(docx_path).stem
        image_dir = os.path.join(output_dir, "images", doc_name)
        
        if DocxDocument is None:
            logger.warning("python-docx not available, skipping DOCX processing.")
            return []

        try:
            doc = DocxDocument(docx_path)
            
            # Extract text content
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            if full_text.strip():
                text_element = ExtractedElement(
                    element_type="text",
                    content=full_text,
                    metadata={
                        "source": docx_path,
                        "document_name": doc_name,
                        "page": 1  # DOCX doesn't have clear page boundaries
                    },
                    page_number=1,
                    element_id=self._generate_element_id(full_text, 1, "text")
                )
                elements.append(text_element)
            
            # Extract images from DOCX
            import zipfile
            
            # Extract images from the DOCX file structure
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for img_index, img_file in enumerate(image_files):
                    try:
                        img_data = docx_zip.read(img_file)
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Save image
                        img_filename = f"img_{img_index}.png"
                        img_path = self._save_image(image, image_dir, img_filename)
                        
                        # Generate description
                        description = self._describe_image_with_vlm(image, full_text)
                        
                        # Determine if it's a graph or regular image
                        is_graph = any(keyword in description.lower() 
                                     for keyword in ['chart', 'graph', 'plot', 'diagram', 'visualization'])
                        element_type = "graph" if is_graph else "image"
                        
                        img_element = ExtractedElement(
                            element_type=element_type,
                            content=img_path,
                            metadata={
                                "source": docx_path,
                                "document_name": doc_name,
                                "image_path": img_path,
                                "description": description,
                                "surrounding_text": full_text,
                                "content_type": element_type,
                                "page": 1
                            },
                            page_number=1,
                            element_id=self._generate_element_id(img_path, 1, element_type)
                        )
                        elements.append(img_element)
                        
                    except Exception as e:
                        logger.error(f"Error processing image {img_index} in DOCX: {e}")
            
            # Extract tables from DOCX
            for table_index, table in enumerate(doc.tables):
                try:
                    # Convert table to HTML
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else [])
                    table_html = df.to_html(index=False, escape=False)
                    
                    # Generate description
                    description = self._extract_table_description(table_html, full_text)
                    
                    table_element = ExtractedElement(
                        element_type="table",
                        content=table_html,
                        metadata={
                            "source": docx_path,
                            "document_name": doc_name,
                            "description": description,
                            "surrounding_text": full_text,
                            "table_shape": df.shape,
                            "content_type": "table",
                            "page": 1
                        },
                        page_number=1,
                        element_id=self._generate_element_id(table_html, 1, "table")
                    )
                    elements.append(table_element)
                    
                except Exception as e:
                    logger.error(f"Error processing table {table_index} in DOCX: {e}")
        
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {e}")
            raise
        
        return elements
    
    def process_document(self, file_path: str, output_dir: str = "extracted_content") -> List[ExtractedElement]:
        """
        Process a document file and extract all content types.
        
        Args:
            file_path: Path to document file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements
        """
        file_extension = Path(file_path).suffix.lower()
        
        # PDF processing
        if file_extension == '.pdf':
            return self.process_pdf(file_path, output_dir)
        
        # Word documents
        elif file_extension in ['.docx', '.doc']:
            return self.process_docx(file_path, output_dir)
        
        # Plain text files
        elif file_extension == '.txt':
            return self.process_txt(file_path, output_dir)
        
        # RTF files
        elif file_extension == '.rtf':
            return self.process_rtf(file_path, output_dir)
        
        # Spreadsheet files (convert to HTML like tables)
        elif file_extension in ['.csv', '.xls', '.xlsx']:
            return self.process_spreadsheet(file_path, output_dir)
        
        # JSON files
        elif file_extension == '.json':
            return self.process_json(file_path, output_dir)
        
        # Markdown files (with hierarchy-aware splitting)
        elif file_extension in ['.md', '.markdown']:
            return self.process_markdown(file_path, output_dir)
        
        # HTML files (similar to markdown)
        elif file_extension in ['.html', '.htm']:
            return self.process_html(file_path, output_dir)
        
        # Image files (process like images in PDF but without surrounding text)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            return self.process_image_file(file_path, output_dir)
        
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .doc, .docx, .txt, .rtf, .csv, .xls, .xlsx, .json, .md, .html, .jpg, .jpeg, .png")
    
    def process_txt(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process TXT file with recursive text splitter and embedding.
        
        Args:
            file_path: Path to TXT file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (text chunks)
        """
        logger.info(f"Processing TXT file: {file_path}")
        
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result.get('encoding', 'utf-8')
        
        # Read text content
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text_content = f.read()
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
                
        if not text_content.strip():
            logger.warning(f"Empty or whitespace-only content in {file_path}")
            return []
        
        # Try to detect page breaks or sections
        page_breaks = self._detect_page_breaks_in_text(text_content)
        
        # Use recursive text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        text_chunks = text_splitter.split_text(text_content)
        
        elements = []
        file_name = Path(file_path).name
        
        for i, chunk in enumerate(text_chunks):
            if chunk.strip():  # Skip empty chunks
                # Estimate page number for this chunk
                estimated_page = self._estimate_page_number(chunk, text_content, page_breaks)
                element_id = f"{file_name}_txt_chunk_{i}"
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=chunk,
                    metadata={
                        "source_file": file_name,
                        "file_type": "txt",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "encoding": encoding,
                        "page_estimated": True if len(page_breaks) > 0 else False
                    },
                    page_number=estimated_page,  # Estimated page number
                    element_id=element_id
                ))
        
        logger.info(f"Extracted {len(elements)} text chunks from {file_path}")
        return elements
    
    def process_rtf(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process RTF file by converting to text and then processing like TXT.
        
        Args:
            file_path: Path to RTF file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (text chunks)
        """
        logger.info(f"Processing RTF file: {file_path}")
        
        if not RTF_AVAILABLE:
            logger.error("striprtf library not available. Cannot process RTF files.")
            raise ImportError("striprtf library required for RTF processing. Install with: pip install striprtf")
        
        # Read RTF file and convert to text
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='cp1252', errors='ignore') as f:
                rtf_content = f.read()
        
        # Convert RTF to plain text
        try:
            text_content = rtf_to_text(rtf_content)
        except Exception as e:
            logger.error(f"Failed to parse RTF content: {e}")
            return []
            
        if not text_content.strip():
            logger.warning(f"Empty or whitespace-only content after RTF conversion in {file_path}")
            return []
        
        # Try to detect page breaks or sections
        page_breaks = self._detect_page_breaks_in_text(text_content)
        
        # Use recursive text splitter (same as TXT processing)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        text_chunks = text_splitter.split_text(text_content)
        
        elements = []
        file_name = Path(file_path).name
        
        for i, chunk in enumerate(text_chunks):
            if chunk.strip():  # Skip empty chunks
                # Estimate page number for this chunk
                estimated_page = self._estimate_page_number(chunk, text_content, page_breaks)
                element_id = f"{file_name}_rtf_chunk_{i}"
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=chunk,
                    metadata={
                        "source_file": file_name,
                        "file_type": "rtf",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "original_format": "rtf",
                        "page_estimated": True if len(page_breaks) > 0 else False
                    },
                    page_number=estimated_page,  # Estimated page number
                    element_id=element_id
                ))
        
        logger.info(f"Extracted {len(elements)} text chunks from RTF file {file_path}")
        return elements
    
    def process_spreadsheet(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process spreadsheet files (CSV, XLS, XLSX) by converting to HTML format like tables.
        
        Args:
            file_path: Path to spreadsheet file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (table elements)
        """
        logger.info(f"Processing spreadsheet file: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        
        if pd is None:
            logger.warning("pandas not available, skipping spreadsheet processing.")
            return []

        try:
            # Read spreadsheet based on format
            if file_extension == '.csv':
                # Try to detect encoding for CSV
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    encoding_result = chardet.detect(raw_data)
                    encoding = encoding_result.get('encoding', 'utf-8')
                
                df = pd.read_csv(file_path, encoding=encoding)
                
            elif file_extension in ['.xls', '.xlsx']:
                # Read Excel file, handle multiple sheets
                excel_file = pd.ExcelFile(file_path)
                
                elements = []
                for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Skip empty sheets
                    if df.empty:
                        logger.warning(f"Empty sheet '{sheet_name}' in {file_path}")
                        continue
                    
                    # Convert DataFrame to HTML
                    html_content = df.to_html(
                        index=False, 
                        classes='spreadsheet-table',
                        table_id=f'sheet_{sheet_idx}',
                        escape=False
                    )
                    
                    # Generate description using LLM (like tables in PDF)
                    description = self._describe_table_with_llm(df, sheet_name)
                    
                    element_id = f"{file_name}_sheet_{sheet_idx}_{sheet_name}"
                    
                    elements.append(ExtractedElement(
                        element_type="table",
                        content=description,  # Store description for embedding
                        metadata={
                            "source_file": file_name,
                            "file_type": file_extension[1:],  # Remove dot
                            "sheet_name": sheet_name,
                            "sheet_index": sheet_idx,
                            "table_html": html_content,
                            "table_shape": df.shape,
                            "columns": df.columns.tolist()
                        },
                        page_number=sheet_idx + 1,  # Use sheet number as page
                        element_id=element_id
                    ))
                
                logger.info(f"Extracted {len(elements)} tables from {len(excel_file.sheet_names)} sheets in {file_path}")
                return elements
            
            else:
                raise ValueError(f"Unsupported spreadsheet format: {file_extension}")
            
            # For CSV files (single sheet)
            if df.empty:
                logger.warning(f"Empty spreadsheet file: {file_path}")
                return []
            
            # Convert DataFrame to HTML
            html_content = df.to_html(
                index=False, 
                classes='spreadsheet-table',
                escape=False
            )
            
            # Generate description using LLM (like tables in PDF)
            description = self._describe_table_with_llm(df, file_name)
            
            element_id = f"{file_name}_table_0"
            
            element = ExtractedElement(
                element_type="table",
                content=description,  # Store description for embedding
                metadata={
                    "source_file": file_name,
                    "file_type": file_extension[1:],  # Remove dot
                    "table_html": html_content,
                    "table_shape": df.shape,
                    "columns": df.columns.tolist(),
                    "encoding": encoding if file_extension == '.csv' else None
                },
                page_number=1,
                element_id=element_id
            )
            
            logger.info(f"Extracted 1 table from spreadsheet {file_path}")
            return [element]
            
        except Exception as e:
            logger.error(f"Error processing spreadsheet {file_path}: {e}")
            return []
    
    def _describe_table_with_llm(self, table_data, table_name: str) -> str:
        """
        Generate a description of the table using LLM.
        
        Args:
            table_data: DataFrame or HTML string containing table data
            table_name: Name/identifier for the table
            
        Returns:
            Text description of the table
        """
        try:
            # Handle different input types
            if isinstance(table_data, str):
                # HTML string - convert to readable text
                table_text = self.html_converter.handle(table_data) if hasattr(self, 'html_converter') else table_data
                table_info = f"Table: {table_name}\n{table_text}"
            elif PANDAS_AVAILABLE and isinstance(table_data, pd.DataFrame):
                # DataFrame - create summary
                table_info = f"Table: {table_name}\n"
                table_info += f"Shape: {table_data.shape[0]} rows, {table_data.shape[1]} columns\n"
                table_info += f"Columns: {', '.join(table_data.columns.tolist())}\n\n"
                
                # Add a sample of the data (first few rows)
                sample_rows = min(5, len(table_data))
                table_info += f"Sample data (first {sample_rows} rows):\n"
                table_info += table_data.head(sample_rows).to_string(index=False)
            else:
                table_info = f"Table: {table_name}\nData: {str(table_data)[:500]}"
            
            # Generate description using LLM
            prompt = f"""Analyze this table and provide a concise description of its content, structure, and key information:

{table_info}

Please provide a brief but informative description that captures:
1. What type of data this table contains
2. Key columns and their purpose

Description:"""

            if hasattr(self, 'llm_model') and self.llm_model:
                try:
                    response = self.llm_model.invoke(prompt)
                    return response.strip() if hasattr(response, 'strip') else str(response).strip()
                except Exception as e:
                    logger.warning(f"LLM description failed for table {table_name}: {e}")
            
            # Fallback: return basic description
            if PANDAS_AVAILABLE and isinstance(table_data, pd.DataFrame):
                return f"Table '{table_name}' with {table_data.shape[0]} rows and {table_data.shape[1]} columns. Columns: {', '.join(table_data.columns.tolist())}. Contains structured data suitable for analysis."
            else:
                return f"Table: {table_name}. Contains structured data in tabular format."
            
        except Exception as e:
            logger.error(f"Error generating table description for {table_name}: {e}")
            return f"Table: {table_name}"
    
    def process_json(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process JSON files by parsing structure and creating text representations.
        
        Args:
            file_path: Path to JSON file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (text or table elements)
        """
        logger.info(f"Processing JSON file: {file_path}")
        
        file_name = Path(file_path).name
        
        try:
            # Read JSON file
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='cp1252', errors='ignore') as f:
                json_data = json.load(f)
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON format in {file_path}: {e}")
            return []
        
        elements = []
        
        # Strategy 1: If JSON is a list of similar objects (like a table)
        if isinstance(json_data, list) and len(json_data) > 0:
            if all(isinstance(item, dict) for item in json_data):
                # Convert list of dicts to DataFrame for table-like processing
                try:
                    df = pd.DataFrame(json_data)
                    
                    # Convert DataFrame to HTML
                    html_content = df.to_html(
                        index=False, 
                        classes='json-table',
                        escape=False
                    )
                    
                    # Generate description using LLM
                    description = self._describe_table_with_llm(df, f"JSON data from {file_name}")
                    
                    element_id = f"{file_name}_json_table_0"
                    
                    elements.append(ExtractedElement(
                        element_type="table",
                        content=description,
                        metadata={
                            "source_file": file_name,
                            "file_type": "json",
                            "data_type": "list_of_objects",
                            "table_html": html_content,
                            "table_shape": df.shape,
                            "columns": df.columns.tolist() if hasattr(df, 'columns') else []
                        },
                        page_number=1,  # JSON files don't have pages
                        element_id=element_id
                    ))
                    
                    logger.info(f"Processed JSON as table with {df.shape[0]} rows and {df.shape[1]} columns")
                    return elements
                    
                except Exception as e:
                    logger.warning(f"Failed to convert JSON list to table: {e}, falling back to text processing")
        
        # Strategy 2: Process as structured text
        # Create a readable text representation of the JSON
        def json_to_readable_text(obj, indent_level=0):
            """Convert JSON object to readable text format."""
            indent = "  " * indent_level
            
            if isinstance(obj, dict):
                if not obj:
                    return "{}"
                text = indent + "{\n"
                for key, value in obj.items():
                    text += f"{indent}  {key}: "
                    if isinstance(value, (dict, list)):
                        text += "\n" + json_to_readable_text(value, indent_level + 1)
                    else:
                        text += str(value)
                    text += "\n"
                text += f"{indent}}}"
                return text
                
            elif isinstance(obj, list):
                if not obj:
                    return "[]"
                text = indent + "[\n"
                for i, item in enumerate(obj):
                    text += f"{indent}  Item {i + 1}: "
                    if isinstance(item, (dict, list)):
                        text += "\n" + json_to_readable_text(item, indent_level + 1)
                    else:
                        text += str(item)
                    text += "\n"
                text += f"{indent}]"
                return text
                
            else:
                return str(obj)
        
        # Convert JSON to readable text
        readable_text = json_to_readable_text(json_data)
        
        # Create a summary description
        summary = self._create_json_summary(json_data, file_name)
        
        # Split the readable text into chunks if it's too long
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,  # Larger chunks for structured data
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", "}", "]", ",", " "]
        )
        
        # Combine summary and readable text
        full_text = f"JSON File Summary:\n{summary}\n\nDetailed Structure:\n{readable_text}"
        text_chunks = text_splitter.split_text(full_text)
        
        for i, chunk in enumerate(text_chunks):
            if chunk.strip():
                element_id = f"{file_name}_json_chunk_{i}"
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=chunk,
                    metadata={
                        "source_file": file_name,
                        "file_type": "json",
                        "data_type": type(json_data).__name__,
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "is_structured_data": True
                    },
                    page_number=1,  # JSON files are single-page conceptually
                    element_id=element_id
                ))
        
        logger.info(f"Processed JSON file into {len(elements)} text chunks")
        return elements
    
    def _create_json_summary(self, json_data, file_name: str) -> str:
        """Create a summary of JSON file structure and content."""
        try:
            summary = f"JSON file: {file_name}\n"
            
            if isinstance(json_data, dict):
                summary += f"Type: Dictionary with {len(json_data)} keys\n"
                summary += f"Keys: {', '.join(list(json_data.keys())[:10])}"  # First 10 keys
                if len(json_data) > 10:
                    summary += " ..."
                    
            elif isinstance(json_data, list):
                summary += f"Type: List with {len(json_data)} items\n"
                if json_data and isinstance(json_data[0], dict):
                    summary += f"List contains dictionaries with keys: {', '.join(list(json_data[0].keys())[:5])}"
                else:
                    summary += f"List contains {type(json_data[0]).__name__} items"
                    
            else:
                summary += f"Type: {type(json_data).__name__}\n"
                summary += f"Value: {str(json_data)[:100]}"
                
            return summary
            
        except Exception as e:
            logger.error(f"Error creating JSON summary: {e}")
            return f"JSON file: {file_name} (summary generation failed)"
    
    def process_markdown(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process Markdown files with hierarchy-aware splitting (# -> ## -> ### -> ... -> \n\n -> \n).
        For images, use surrounding text (text before and after), preserving order.
        
        Args:
            file_path: Path to Markdown file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (text chunks and images)
        """
        logger.info(f"Processing Markdown file: {file_path}")
        
        # Read Markdown file
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                markdown_content = f.read()
                
        if not markdown_content.strip():
            logger.warning(f"Empty or whitespace-only content in {file_path}")
            return []
        
        elements = []
        file_name = Path(file_path).name
        
        # Extract images from markdown first
        import re
        image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        images_found = list(re.finditer(image_pattern, markdown_content))
        
        # Process images with surrounding text
        for img_match in images_found:
            alt_text = img_match.group(1)
            img_src = img_match.group(2)
            img_start = img_match.start()
            img_end = img_match.end()
            
            # Get surrounding text (before and after the image)
            context_range = 200  # Characters before and after
            text_before = markdown_content[max(0, img_start - context_range):img_start].strip()
            text_after = markdown_content[img_end:min(len(markdown_content), img_end + context_range)].strip()
            surrounding_text = f"{text_before} [IMAGE: {alt_text}] {text_after}".strip()
            
            # Try to process the image if it's a local file
            try:
                # Handle relative paths and local images
                if not img_src.startswith(('http://', 'https://', 'ftp://')):
                    # Local image file
                    img_full_path = os.path.join(os.path.dirname(file_path), img_src)
                    if os.path.exists(img_full_path):
                        image = Image.open(img_full_path)
                        
                        # Save image to output directory
                        image_dir = os.path.join(output_dir, "images")
                        os.makedirs(image_dir, exist_ok=True)
                        image_filename = f"{file_name}_{len(elements)}_{Path(img_src).name}"
                        image_path = os.path.join(image_dir, image_filename)
                        image.save(image_path)
                        
                        # Generate description with surrounding text
                        description = self._describe_image_with_vlm(image, surrounding_text)
                        
                        # Determine if it's a graph or regular image
                        is_graph = any(keyword in description.lower() 
                                     for keyword in ['chart', 'graph', 'plot', 'diagram', 'visualization'])
                        element_type = "graph" if is_graph else "image"
                        
                        img_element = ExtractedElement(
                            element_type=element_type,
                            content=description,  # Store description for embedding
                            metadata={
                                "source_file": file_name,
                                "file_type": "md",
                                "image_path": image_path,
                                "alt_text": alt_text,
                                "original_src": img_src,
                                "surrounding_text": surrounding_text,
                                "markdown_context": True
                            },
                            page_number=1,  # Markdown files are typically single-page
                            element_id=f"{file_name}_img_{len(elements)}"
                        )
                        elements.append(img_element)
                        
            except Exception as e:
                logger.warning(f"Could not process image {img_src} in markdown: {e}")
        
        # Remove image markdown syntax for text processing
        text_content = re.sub(image_pattern, '', markdown_content)
        
        # Try to detect page/section breaks in markdown
        page_breaks = self._detect_page_breaks_in_text(text_content)
        
        # Use MarkdownTextSplitter with hierarchy awareness
        if markdown is not None:
            # Use markdown-aware splitter
            markdown_splitter = MarkdownTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
        else:
            # Fallback to regular text splitter with markdown-like separators
            markdown_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=[
                    "\n# ",      # H1
                    "\n## ",     # H2
                    "\n### ",    # H3
                    "\n#### ",   # H4
                    "\n##### ",  # H5
                    "\n###### ", # H6
                    "\n\n",      # Paragraphs
                    "\n",        # Lines
                    " ",         # Words
                    ""           # Characters
                ]
            )
        
        text_chunks = markdown_splitter.split_text(text_content)
        
        for i, chunk in enumerate(text_chunks):
            if chunk.strip():  # Skip empty chunks
                # Estimate page/section number for this chunk
                estimated_page = self._estimate_page_number(chunk, text_content, page_breaks)
                element_id = f"{file_name}_md_chunk_{i}"
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=chunk,
                    metadata={
                        "source_file": file_name,
                        "file_type": "md",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "is_markdown": True,
                        "page_estimated": True if len(page_breaks) > 0 else False
                    },
                    page_number=estimated_page,  # Estimated page/section number
                    element_id=element_id
                ))
        
        logger.info(f"Extracted {len(elements)} elements from Markdown file {file_path}")
        return elements
    
    def process_html(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process HTML files by parsing content and extracting text and images (similar to markdown).
        
        Args:
            file_path: Path to HTML file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (text chunks and images)
        """
        logger.info(f"Processing HTML file: {file_path}")
        
        if not BS4_AVAILABLE:
            logger.warning("BeautifulSoup not available, using basic HTML processing.")
            # Fallback to basic text extraction
            return self._process_html_basic(file_path, output_dir)
        
        # Read HTML file
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
                
        if not html_content.strip():
            logger.warning(f"Empty or whitespace-only content in {file_path}")
            return []
        
        elements = []
        file_name = Path(file_path).name
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract images from HTML
        images = soup.find_all('img')
        for img_idx, img_tag in enumerate(images):
            img_src = img_tag.get('src', '')
            alt_text = img_tag.get('alt', '')
            
            # Get surrounding text (previous and next siblings)
            context_elements = []
            
            # Get previous siblings (up to 2)
            prev_siblings = []
            for sibling in img_tag.previous_siblings:
                if hasattr(sibling, 'get_text'):
                    prev_siblings.append(sibling.get_text().strip())
                elif isinstance(sibling, str):
                    prev_siblings.append(sibling.strip())
                if len(prev_siblings) >= 2:
                    break
            
            # Get next siblings (up to 2) 
            next_siblings = []
            for sibling in img_tag.next_siblings:
                if hasattr(sibling, 'get_text'):
                    next_siblings.append(sibling.get_text().strip())
                elif isinstance(sibling, str):
                    next_siblings.append(sibling.strip())
                if len(next_siblings) >= 2:
                    break
            
            surrounding_text = ' '.join(reversed(prev_siblings)) + f" [IMAGE: {alt_text}] " + ' '.join(next_siblings)
            surrounding_text = surrounding_text.strip()
            
            # Try to process the image if it's a local file
            try:
                if not img_src.startswith(('http://', 'https://', 'ftp://', 'data:')):
                    # Local image file
                    img_full_path = os.path.join(os.path.dirname(file_path), img_src)
                    if os.path.exists(img_full_path):
                        image = Image.open(img_full_path)
                        
                        # Save image to output directory
                        image_dir = os.path.join(output_dir, "images")
                        os.makedirs(image_dir, exist_ok=True)
                        image_filename = f"{file_name}_{img_idx}_{Path(img_src).name}"
                        image_path = os.path.join(image_dir, image_filename)
                        image.save(image_path)
                        
                        # Generate description with surrounding text
                        description = self._describe_image_with_vlm(image, surrounding_text)
                        
                        # Determine if it's a graph or regular image
                        is_graph = any(keyword in description.lower() 
                                     for keyword in ['chart', 'graph', 'plot', 'diagram', 'visualization'])
                        element_type = "graph" if is_graph else "image"
                        
                        img_element = ExtractedElement(
                            element_type=element_type,
                            content=description,  # Store description for embedding
                            metadata={
                                "source_file": file_name,
                                "file_type": "html",
                                "image_path": image_path,
                                "alt_text": alt_text,
                                "original_src": img_src,
                                "surrounding_text": surrounding_text,
                                "html_context": True
                            },
                            page_number=1,
                            element_id=f"{file_name}_img_{img_idx}"
                        )
                        elements.append(img_element)
                        
            except Exception as e:
                logger.warning(f"Could not process image {img_src} in HTML: {e}")
        
        # Convert HTML to text while preserving structure
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True  # We already processed images separately
        h.body_width = 0  # Don't wrap lines
        
        text_content = h.handle(html_content)
        
        # Try to detect page breaks or sections
        page_breaks = self._detect_page_breaks_in_text(text_content)
        
        # Use recursive text splitter with HTML-aware separators
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=[
                "\n# ",      # H1 (from html2text conversion)
                "\n## ",     # H2
                "\n### ",    # H3
                "\n#### ",   # H4
                "\n##### ",  # H5
                "\n###### ", # H6
                "\n\n",      # Paragraphs
                "\n",        # Lines
                " ",         # Words
                ""           # Characters
            ]
        )
        
        text_chunks = text_splitter.split_text(text_content)
        
        for i, chunk in enumerate(text_chunks):
            if chunk.strip():  # Skip empty chunks
                element_id = f"{file_name}_html_chunk_{i}"
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=chunk,
                    metadata={
                        "source_file": file_name,
                        "file_type": "html",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "is_html": True,
                        "page_estimated": True if len(page_breaks) > 0 else False
                    },
                    page_number=self._estimate_page_number(chunk, text_content, page_breaks), # Estimated page number
                    element_id=element_id
                ))
        
        logger.info(f"Extracted {len(elements)} elements from HTML file {file_path}")
        return elements
    
    def _process_html_basic(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Basic HTML processing without BeautifulSoup (fallback method).
        """
        logger.info(f"Processing HTML file with basic method: {file_path}")
        
        # Read HTML file
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
        
        # Convert HTML to text using html2text
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True
        h.body_width = 0
        
        text_content = h.handle(html_content)
        
        # Try to detect page breaks
        page_breaks = self._detect_page_breaks_in_text(text_content)
        
        # Use recursive text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        
        text_chunks = text_splitter.split_text(text_content)
        
        elements = []
        file_name = Path(file_path).name
        
        for i, chunk in enumerate(text_chunks):
            if chunk.strip():
                estimated_page = self._estimate_page_number(chunk, text_content, page_breaks)
                element_id = f"{file_name}_html_basic_chunk_{i}"
                
                elements.append(ExtractedElement(
                    element_type="text",
                    content=chunk,
                    metadata={
                        "source_file": file_name,
                        "file_type": "html",
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "is_html": True,
                        "basic_processing": True,
                        "page_estimated": True if len(page_breaks) > 0 else False
                    },
                    page_number=estimated_page,
                    element_id=element_id
                ))
        
        logger.info(f"Extracted {len(elements)} elements from HTML file using basic processing")
        return elements
    
    def process_image_file(self, file_path: str, output_dir: str) -> List[ExtractedElement]:
        """
        Process standalone image files (JPEG, PNG, etc.) like images in PDF but without surrounding text.
        
        Args:
            file_path: Path to image file
            output_dir: Directory to save extracted content
            
        Returns:
            List of extracted elements (single image element)
        """
        logger.info(f"Processing standalone image file: {file_path}")
        
        try:
            # Open and validate image
            image = Image.open(file_path)
            
            # Save image to output directory
            image_dir = os.path.join(output_dir, "images")
            os.makedirs(image_dir, exist_ok=True)
            image_filename = Path(file_path).name
            image_output_path = os.path.join(image_dir, image_filename)
            
            # Save in a common format if needed
            if image.mode in ('RGBA', 'LA', 'P'):
                # Convert to RGB for better compatibility
                rgb_image = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                rgb_image.paste(image, mask=image.split()[-1] if image.mode in ('RGBA', 'LA') else None)
                image = rgb_image
            
            image.save(image_output_path)
            
            # Generate description without surrounding text (as specified)
            surrounding_text = ""  # No surrounding text for standalone images
            description = self._describe_image_with_vlm(image, surrounding_text)
            
            # Determine if it's a graph or regular image based on content
            is_graph = any(keyword in description.lower() 
                         for keyword in ['chart', 'graph', 'plot', 'diagram', 'visualization', 
                                       'data', 'statistics', 'analytics', 'metrics'])
            element_type = "graph" if is_graph else "image"
            
            file_name = Path(file_path).name
            element_id = self._generate_element_id(file_name, 1, element_type)
            
            img_element = ExtractedElement(
                element_type=element_type,
                content=description,  # Store description for embedding
                metadata={
                    "source_file": file_name,
                    "file_type": Path(file_path).suffix[1:].lower(),  # Extension without dot
                    "image_path": image_output_path,
                    "original_path": file_path,
                    "image_size": image.size,
                    "image_mode": image.mode,
                    "standalone_image": True,
                    "surrounding_text": surrounding_text  # Empty for standalone images
                },
                page_number=1,
                element_id=element_id,
                bbox=None  # No bounding box for standalone images
            )
            
            logger.info(f"Processed standalone image {file_path} as {element_type}")
            return [img_element]
            
        except Exception as e:
            logger.error(f"Error processing standalone image file {file_path}: {e}")
            return []
    
    def _extract_text_with_markdown_strategy(self, pdf_path: str) -> List[ExtractedElement]:
        """
        Extract text from PDF using pymupdf4llm to get markdown format, 
        then process with MarkdownTextSplitter for better chunking.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of text elements with markdown-aware chunking
        """
        elements = []
        
        if not PYMUPDF4LLM_AVAILABLE:
            logger.warning("pymupdf4llm not available, falling back to regular text extraction")
            return self._extract_text_elements(pdf_path)
        
        try:
            logger.info(f"Extracting text as markdown from PDF: {pdf_path}")
            
            # Extract markdown text using pymupdf4llm with page chunks to preserve page info
            try:
                # Try to get page chunks to preserve page information
                markdown_data = pymupdf4llm.to_markdown(
                    pdf_path,
                    pages=None,  # Process all pages
                    page_chunks=True,  # Get page-level chunks with metadata
                    write_images=False,  # We handle images separately
                    embed_images=False  # Don't embed images in text
                )
                
                # If we get page chunks, process them individually
                if isinstance(markdown_data, list) and len(markdown_data) > 0:
                    logger.info(f"Got {len(markdown_data)} page chunks from pymupdf4llm")
                    
                    for page_data in markdown_data:
                        if isinstance(page_data, dict) and 'text' in page_data:
                            page_text = page_data['text']
                            page_metadata = page_data.get('metadata', {})
                            page_number = page_metadata.get('page', 1)  # Get real page number
                            
                            if page_text and page_text.strip():
                                # Use MarkdownTextSplitter for this page's content
                                markdown_splitter = MarkdownTextSplitter(
                                    chunk_size=self.chunk_size,
                                    chunk_overlap=self.chunk_overlap,
                                    length_function=len,
                                    is_separator_regex=False,
                                )
                                
                                chunks = markdown_splitter.split_text(page_text)
                                
                                for i, chunk_text in enumerate(chunks):
                                    if chunk_text.strip():
                                        element = ExtractedElement(
                                            element_type="text",
                                            content=chunk_text.strip(),
                                            metadata={
                                                "source": pdf_path,
                                                "page": page_number,  # Real page number
                                                "chunk_index": i,
                                                "total_chunks_in_page": len(chunks),
                                                "extraction_method": "pymupdf4llm_markdown_pages",
                                                "text_splitter": "MarkdownTextSplitter",
                                                "page_metadata": page_metadata
                                            },
                                            page_number=page_number,  # Real page number
                                            element_id=self._generate_element_id(chunk_text, page_number, f"text_p{page_number}_c{i}")
                                        )
                                        elements.append(element)
                    
                    logger.info(f"Created {len(elements)} text elements using page-aware markdown strategy")
                    return elements
                    
            except Exception as e:
                logger.warning(f"Page chunks extraction failed: {e}, trying regular markdown extraction")
            
            # Fallback: Extract as single markdown text (original method)
            markdown_text = pymupdf4llm.to_markdown(
                pdf_path,
                pages=None,  # Process all pages
                write_images=False,  # We handle images separately
                embed_images=False  # Don't embed images in text
            )
            
            if not markdown_text or not markdown_text.strip():
                logger.warning("No markdown text extracted, falling back to regular extraction")
                return self._extract_text_elements(pdf_path)
            
            logger.info(f"Extracted {len(markdown_text)} characters of markdown text")
            
            # Try to detect page breaks in the markdown text to assign better page numbers
            page_breaks = self._detect_page_breaks_in_text(markdown_text)
            
            # Use MarkdownTextSplitter for hierarchy-aware chunking
            markdown_splitter = MarkdownTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                is_separator_regex=False,
            )
            
            # Split the markdown text into chunks
            chunks = markdown_splitter.split_text(markdown_text)
            
            logger.info(f"Split markdown text into {len(chunks)} chunks")
            
            # Create ExtractedElement for each chunk with estimated page numbers
            for i, chunk_text in enumerate(chunks):
                if chunk_text.strip():  # Only add non-empty chunks
                    # Estimate page number based on chunk position and page breaks
                    estimated_page = self._estimate_page_number(chunk_text, markdown_text, page_breaks)
                    
                    element = ExtractedElement(
                        element_type="text",
                        content=chunk_text.strip(),
                        metadata={
                            "source": pdf_path,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "extraction_method": "pymupdf4llm_markdown",
                            "text_splitter": "MarkdownTextSplitter",
                            "page_estimated": True
                        },
                        page_number=estimated_page,  # Estimated page number
                        element_id=self._generate_element_id(chunk_text, estimated_page, f"text_est{estimated_page}_c{i}")
                    )
                    elements.append(element)
            
            logger.info(f"Created {len(elements)} text elements using markdown strategy")
            return elements
            
        except Exception as e:
            logger.error(f"Error in markdown text extraction: {e}")
            logger.info("Falling back to regular text extraction")
            return self._extract_text_elements(pdf_path)
    
    def _detect_page_breaks_in_text(self, text: str) -> List[int]:
        """
        Detect potential page breaks in text to help estimate page numbers.
        
        Args:
            text: Full text content
            
        Returns:
            List of character positions where page breaks likely occur
        """
        import re
        
        page_breaks = []
        
        # Look for common page break indicators
        patterns = [
            r'\n\s*Page\s+\d+\s*\n',  # "Page X" indicators
            r'\n\s*\d+\s*\n',         # Standalone page numbers
            r'\f',                     # Form feed characters
            r'\n\s*[-=]{3,}\s*\n',    # Horizontal lines
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                page_breaks.append(match.start())
        
        # Sort and remove duplicates
        page_breaks = sorted(list(set(page_breaks)))
        
        # If no explicit breaks found, estimate based on text length
        if not page_breaks:
            # Estimate page breaks every ~2000-3000 characters (rough page estimate)
            avg_chars_per_page = 2500
            for i in range(avg_chars_per_page, len(text), avg_chars_per_page):
                page_breaks.append(i)
        
        return page_breaks
    
    def _estimate_page_number(self, chunk_text: str, full_text: str, page_breaks: List[int]) -> int:
        """
        Estimate the page number for a text chunk.
        
        Args:
            chunk_text: The text chunk to find page for
            full_text: Full document text
            page_breaks: List of character positions where pages likely break
            
        Returns:
            Estimated page number
        """
        try:
            # Find the position of this chunk in the full text
            chunk_position = full_text.find(chunk_text)
            
            if chunk_position == -1:
                return 1  # Fallback if chunk not found
            
            # Count how many page breaks occur before this position
            page_number = 1
            for break_pos in page_breaks:
                if break_pos < chunk_position:  # Breaks that occur before this chunk
                    page_number += 1
                else:
                    break
            
            return page_number
            
        except Exception as e:
            logger.warning(f"Error estimating page number: {e}")
            return 1
    
    def _extract_text_elements(self, pdf_path: str) -> List[ExtractedElement]:
        """
        Extract text elements from PDF using regular PyMuPDF text extraction.
        This is the fallback method when pymupdf4llm is not available.
        Now extracts real page numbers when possible.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of text elements
        """
        elements = []
        
        try:
            doc = fitz.open(pdf_path)
            
            # Process each page individually to preserve page numbers
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    # Use RecursiveCharacterTextSplitter for this page's text
                    chunks = self.text_splitter.split_text(page_text)
                    
                    # Create ExtractedElement for each chunk with real page number
                    for i, chunk_text in enumerate(chunks):
                        if chunk_text.strip():
                            element = ExtractedElement(
                                element_type="text",
                                content=chunk_text.strip(),
                                metadata={
                                    "source": pdf_path,
                                    "page": page_num + 1,  # Real page number (1-indexed)
                                    "chunk_index_in_page": i,
                                    "total_chunks_in_page": len(chunks),
                                    "extraction_method": "pymupdf_regular_page_aware",
                                    "text_splitter": "RecursiveCharacterTextSplitter"
                                },
                                page_number=page_num + 1,  # Real page number (1-indexed)
                                element_id=self._generate_element_id(chunk_text, page_num + 1, f"text_p{page_num + 1}_c{i}")
                            )
                            elements.append(element)
            
            doc.close()
            
            logger.info(f"Extracted {len(elements)} text chunks using page-aware regular strategy")
            return elements
            
        except Exception as e:
            logger.error(f"Error in regular text extraction: {e}")
            return elements
    
    def elements_to_documents(self, elements: List[ExtractedElement]) -> List[Document]:
        """
        Convert extracted elements to Langchain Documents based on their type.
        
        Args:
            elements: List of extracted elements
            
        Returns:
            List of Langchain Documents ready for embedding
        """
        documents = []
        
        for element in elements:
            if element.element_type == "text":
                # Split text into chunks
                text_docs = self.text_splitter.create_documents(
                    [element.content],
                    metadatas=[element.metadata]
                )
                
                # Add unique IDs to text chunks
                for i, doc in enumerate(text_docs):
                    doc.metadata["id"] = f"{element.element_id}_chunk_{i}"
                    doc.metadata["element_type"] = "text"
                    doc.metadata["original_element_id"] = element.element_id
                
                documents.extend(text_docs)
                
            elif element.element_type in ["image", "graph"]:
                # For images and graphs, embed the description but keep image path for retrieval
                doc = Document(
                    page_content=element.metadata.get("description", ""),
                    metadata={
                        **element.metadata,
                        "id": element.element_id,
                        "element_type": element.element_type,
                        "image_path": element.content,
                        "retrieve_content": element.content  # Store image path for retrieval
                    }
                )
                documents.append(doc)
                
            elif element.element_type == "table":
                # For tables, embed the description but keep HTML structure
                doc = Document(
                    page_content=element.metadata.get("description", ""),
                    metadata={
                        **element.metadata,
                        "id": element.element_id,
                        "element_type": "table",
                        "table_html": element.content,
                        "retrieve_content": element.content  # Store HTML for retrieval
                    }
                )
                documents.append(doc)
                
            elif element.element_type == "page_image":
                # Page images will be handled separately with multimodal embeddings
                doc = Document(
                    page_content=f"Page {element.page_number} image",
                    metadata={
                        **element.metadata,
                        "id": element.element_id,
                        "element_type": "page_image",
                        "image_path": element.content,
                        "retrieve_content": element.content
                    }
                )
                documents.append(doc)
        
        return documents

if __name__ == "__main__":
    # Example usage
    processor = MultimodalDocumentProcessor()
    
    # Process a document
    file_path = "example.pdf"  # Replace with actual file path
    elements = processor.process_document(file_path)
    
    print(f"Extracted {len(elements)} elements:")
    for element in elements:
        print(f"- {element.element_type}: {element.element_id}")
    
    # Convert to documents
    documents = processor.elements_to_documents(elements)
    print(f"Created {len(documents)} documents for embedding") 